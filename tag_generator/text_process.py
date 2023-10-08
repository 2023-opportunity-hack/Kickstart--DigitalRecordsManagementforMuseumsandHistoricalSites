def text_process(file_path):
    import torch
    # import SentencePiece
    # extract text from PDF
    import PyPDF2

    from transformers import (
        # headline generation libraries
        T5ForConditionalGeneration,
        T5Tokenizer,

        # extract significant text libraries
        TokenClassificationPipeline,
        AutoModelForTokenClassification,
        AutoTokenizer,
    )

    # for significant text tag generation
    from transformers.pipelines import AggregationStrategy
    import numpy as np

    # from pptx import Presentation
    import docx

    # PDF TEXT EXTRACTION
    def extract_text_from_pdf(pdf_file_path):
        text = ""
        try:
            with open(pdf_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    # page = pdf_reader.getPage(page_num)
                    text += page.extract_text()
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
        return text

    def extract_text_from_txt_file(file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text
        except FileNotFoundError:
            return f"File not found: {file_path}"
        except Exception as e:
            return f"An error occurred: {str(e)}"

    # def extract_text_from_powerpoint(file_path):
    #     # Replace 'your_presentation.pptx' with the actual path to your PowerPoint file
    #     # pptx_file = 'your_presentation.pptx'

    #     # Load the PowerPoint presentation
    #     presentation = Presentation(file_path, encoding='ISO-8859-1')

    #     all_text = ""

    #     # Iterate through slides and extract text content
    #     for slide in presentation.slides:
    #         for shape in slide.shapes:
    #             if shape.has_text_frame:
    #                 text_frame = shape.text_frame
    #                 for paragraph in text_frame.paragraphs:
    #                     for run in paragraph.runs:
    #                         try:
    #                             text = run.text.encode('ISO-8859-1', 'ignore').decode('ISO-8859-1')
    #                             all_text += text + "\n"  # Add a newline separator
    #                         except UnicodeEncodeError:
    #                             all_text += "Encoding Error\n"

    #                             # Close the presentation file
    #     presentation.close()

    #     # Print or use the 'all_text' variable
    #     return all_text

    def extract_from_docx(file_path):
        # Replace 'your_document.docx' with the actual path to your Word document
        docx_file = 'your_document.docx'

        # Load the Word document
        doc = docx.Document(file_path)

        # Initialize an empty string to store the extracted text
        all_text = ""

        # Iterate through paragraphs and extract text
        for paragraph in doc.paragraphs:
            text = paragraph.text
            all_text += text + "\n"  # Add a newline separator
        return all_text

    # HEADLINE GENERATION
    def headline_generation(article):
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model = T5ForConditionalGeneration.from_pretrained("Michau/t5-base-en-generate-headline")
        tokenizer = T5Tokenizer.from_pretrained("Michau/t5-base-en-generate-headline")
        model = model.to(device)
        max_len = 256
        encoding = tokenizer.encode_plus(article, return_tensors="pt")
        input_ids = encoding["input_ids"].to(device)
        attention_masks = encoding["attention_mask"].to(device)
        beam_outputs = model.generate(
            input_ids=input_ids,
            attention_mask=attention_masks,
            max_length=64,
            num_beams=3,
            early_stopping=True,
        )

        headline = tokenizer.decode(beam_outputs[0])
        import re
        pattern = r'<[^>]+>'
        result = re.sub(pattern, '', headline)
        return result.lstrip()

    # KEYPHRASE EXTRACTION
    # Define keyphrase extraction pipeline
    def KeyphraseExtration(article):
        # class KeyphraseExtractionPipeline(TokenClassificationPipeline):
        #     def __init__(self, model, *args, **kwargs):
        #         super().__init__(
        #             model=AutoModelForTokenClassification.from_pretrained(model),
        #             tokenizer=AutoTokenizer.from_pretrained(model),
        #             *args,
        #             **kwargs
        #         )
        #
        #     def postprocess(self, all_outputs):
        #         results = super().postprocess(
        #             all_outputs=all_outputs,
        #             aggregation_strategy=AggregationStrategy.FIRST,
        #         )
        #         return np.unique([result.get("word").strip() for result in results])
        #
        # # Load pipeline
        # model_name = "ml6team/keyphrase-extraction-distilbert-inspec"
        # extractor = KeyphraseExtractionPipeline(model=model_name)
        # keyphrases = extractor(article)
        # ans = []
        # for x in keyphrases:
        #     ans.append(x)
        # return ans



        import spacy
        from collections import Counter

        # Load the English language model
        nlp = spacy.load("en_core_web_sm")

        # # Read the text file
        # with open(article, "r", encoding="utf-8") as file:
        #     text = file.read()

        # Process the text using spaCy
        doc = nlp(article)

        # Filter out stopwords, punctuation marks, and '\n'
        keywords = [token.text.lower() for token in doc if
                    not token.is_stop and not token.is_punct and token.text != '\n']

        # Count the frequency of each keyword
        keyword_freq = Counter(keywords)

        # Get the 10 most common keywords
        most_common_keywords = keyword_freq.most_common(10)

        # Extract the words from the tuples and store them in a list
        top_10_words = [keyword for keyword, freq in most_common_keywords]

        # Return the list of the 10 most common words
        return top_10_words

    if file_path.endswith(".pdf"):
        article_text = "headline: " + extract_text_from_pdf(file_path)
    elif file_path.endswith(".txt"):
        article_text = "headline: " + extract_text_from_txt_file(file_path)
    # elif file_path.endswith(".pptx"):
    #     article_text = "headline: " + extract_text_from_txt_file(file_path)
    elif file_path.endswith(".docx"):
        article_text = "headline: " + extract_from_docx(file_path)
    # elif file_path.endswith(".doc"):
    #     article_text = "headline: " + extract_from_doc(file_path)
    # elif file_path.endswith(".wpd"):
    #     article_text = "headline: " + extract_from_wpd(file_path)
    else:
        raise TypeError("File Type Not Recognized")

    headline = headline_generation(article_text)
    keyphrases = KeyphraseExtration(article_text)
    # PRINT HEADLINE AND KEYPHRASES
    print("Headline: ", headline)
    print("Key Phrases: ", keyphrases)
    return [headline] + keyphrases
